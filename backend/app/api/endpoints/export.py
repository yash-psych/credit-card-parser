import io
import json # Import the json library
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
import pandas as pd
from docx import Document
from fpdf import FPDF
from typing import Optional
from datetime import datetime, timedelta

from app.api import deps
from app.models import models

router = APIRouter()

def get_user_data_as_dataframe(db: Session, user: models.User, issuer: Optional[str] = None, period: Optional[str] = None):
    """
    Fetches user's upload history with filters and returns it as a pandas DataFrame.
    """
    query = db.query(models.FileUpload).filter(models.FileUpload.user_id == user.id)

    if issuer:
        query = query.filter(models.FileUpload.issuer == issuer)
    
    if period:
        today = datetime.utcnow()
        if period == "day": start_date = today - timedelta(days=1)
        elif period == "week": start_date = today - timedelta(weeks=1)
        elif period == "month": start_date = today - timedelta(days=30)
        elif period == "year": start_date = today - timedelta(days=365)
        else: start_date = None
        
        if start_date:
            query = query.filter(models.FileUpload.uploaded_at >= start_date)

    uploads = query.order_by(models.FileUpload.uploaded_at.desc()).all()
    
    records = []
    for upload in uploads:
        # --- THIS IS THE FIX ---
        # Ensure the extracted_data string is parsed back into a dictionary
        if isinstance(upload.extracted_data, str):
            data = json.loads(upload.extracted_data)
        else:
            data = upload.extracted_data or {}
        # --- END FIX ---

        records.append({
            'Filename': upload.filename,
            'Issuer': upload.issuer,
            'Card (Last 4)': data.get('last_4_digits'), 'Card Variant': data.get('card_variant'),
            'Billing Cycle': data.get('billing_cycle'), 'Payment Due Date': data.get('payment_due_date'),
            'Total Balance': data.get('total_balance')
        })
    return pd.DataFrame(records)

@router.get("/export/xlsx")
async def export_to_xlsx(
    issuer: Optional[str] = Query(None),
    period: Optional[str] = Query(None),
    db: Session = Depends(deps.get_db),
    current_user: models.User = Depends(deps.get_current_user_from_token)
):
    df = get_user_data_as_dataframe(db, current_user, issuer, period)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Extracted Data')
    output.seek(0)
    headers = {'Content-Disposition': 'attachment; filename="exported_data.xlsx"'}
    return StreamingResponse(output, headers=headers, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@router.get("/export/pdf")
async def export_to_pdf(
    issuer: Optional[str] = Query(None),
    period: Optional[str] = Query(None),
    db: Session = Depends(deps.get_db),
    current_user: models.User = Depends(deps.get_current_user_from_token)
):
    df = get_user_data_as_dataframe(db, current_user, issuer, period)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Extracted Credit Card Data", ln=True, align='C')
    for i, row in df.iterrows():
        pdf.cell(200, 10, txt=f"--- Record {i+1} ---", ln=True)
        for col, value in row.items():
            pdf.cell(200, 10, txt=f"{col}: {value}", ln=True)
        pdf.cell(200, 5, txt="", ln=True)
    output = io.BytesIO(pdf.output(dest='S').encode('latin-1'))
    headers = {'Content-Disposition': 'attachment; filename="exported_data.pdf"'}
    return StreamingResponse(output, headers=headers, media_type='application/pdf')

@router.get("/export/docx")
async def export_to_docx(
    issuer: Optional[str] = Query(None),
    period: Optional[str] = Query(None),
    db: Session = Depends(deps.get_db),
    current_user: models.User = Depends(deps.get_current_user_from_token)
):
    df = get_user_data_as_dataframe(db, current_user, issuer, period)
    document = Document()
    document.add_heading('Extracted Credit Card Data', 0)
    for i, row in df.iterrows():
        document.add_heading(f"Record {i+1}: {row['Filename']}", level=2)
        for col, value in row.items():
            if col != 'Filename': document.add_paragraph(f"{col}: {value}")
        document.add_paragraph()
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    headers = {'Content-Disposition': 'attachment; filename="exported_data.docx"'}
    return StreamingResponse(output, headers=headers, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')